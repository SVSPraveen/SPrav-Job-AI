import os
from docxtpl import DocxTemplate
from docx import Document

def create_default_template(template_path: str):
    """Generates a default Jinja-tagged Word template for docxtpl to consume."""
    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    doc = Document()
    doc.add_heading("{{ personal.name }}", 0)
    
    p = doc.add_paragraph()
    p.add_run("{{ personal.email }} | {{ personal.phone }} | {{ personal.linkedin }} | {{ personal.github }}")
    
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("{{ summary }}")
    
    doc.add_heading("Experience", level=1)
    
    # We use jinja loops to iterate over hydrated bullets, grouped by parent_id if needed.
    # For a basic template, we'll just list the bullets.
    # Ideally, we structure the context for the template to group bullets by job.
    doc.add_paragraph("{% for exp in experiences %}")
    doc.add_heading("{{ exp.company }} - {{ exp.role }}", level=2)
    doc.add_paragraph("{{ exp.start_date }} - {{ exp.end_date }}")
    doc.add_paragraph("{% for bullet in exp.bullets %}")
    doc.add_paragraph("• {{ bullet.text }}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endfor %}")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("Languages: {{ skills.languages | join(', ') }}")
    doc.add_paragraph("Frameworks: {{ skills.frameworks | join(', ') }}")
    doc.add_paragraph("Tools: {{ skills.tools | join(', ') }}")

    doc.save(template_path)

def generate_docx(tailored_resume: dict, kb: dict, output_path: str, template_path: str = "templates/resume_template.docx"):
    """
    Renders the tailored resume into a .docx file using a template.
    """
    if not os.path.exists(template_path):
        create_default_template(template_path)

    # Prepare context for the template
    context = {
        "personal": kb.get("personal", {}),
        "summary": tailored_resume.get("summary", ""),
        "skills": kb.get("skills", {}),
        "experiences": []
    }

    # Group hydrated bullets by their parent (work_history or project)
    hydrated_bullets = tailored_resume.get("hydrated_bullets", [])
    work_histories = {w["id"]: w for w in kb.get("work_history", [])}
    projects = {p["id"]: p for p in kb.get("projects", [])}

    exp_map = {}
    for bullet in hydrated_bullets:
        pid = bullet.get("parent_id")
        if pid not in exp_map:
            # Figure out if it's a job or project
            parent = work_histories.get(pid) or projects.get(pid)
            if parent:
                exp_map[pid] = {
                    "company": parent.get("company") or parent.get("name"),
                    "role": parent.get("role") or parent.get("tagline"),
                    "start_date": parent.get("start_date", ""),
                    "end_date": parent.get("end_date", ""),
                    "bullets": []
                }
        if pid in exp_map:
            exp_map[pid]["bullets"].append(bullet)

    context["experiences"] = list(exp_map.values())

    # Render template
    tpl = DocxTemplate(template_path)
    tpl.render(context)
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    tpl.save(output_path)
    
    # Metadata update using python-docx on the saved file
    try:
        from docx import Document
        doc = Document(output_path)
        author_name = kb.get("personal", {}).get("name", "AutoJob AI")
        core_props = doc.core_properties
        core_props.author = author_name
        core_props.last_modified_by = author_name
        core_props.comments = ""
        doc.save(output_path)
    except Exception as e:
        print(f"[Formatter] Failed to update DOCX metadata: {e}")
    return output_path

def generate_pdf(tailored_resume: dict, kb: dict, pdf_path: str) -> bool:
    """
    Generates a 1-page PDF using ReportLab with dynamic font sizing and relevance-based truncation.
    """
    import json
    print("\n--- RAW DATA RECEIVED BY FORMATTER.PY ---")
    print(json.dumps(tailored_resume, indent=2))
    print("-------------------------------------------\n")

    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    
    personal = kb.get("personal", {})
    name = personal.get("name", "")
    # Contact line: location leads, matching the reference resume's format exactly.
    contact_parts = [
        personal.get("location", ""),
        personal.get("phone", ""),
        personal.get("email", ""),
        personal.get("linkedin", ""),
        personal.get("github", ""),
    ]
    contact_info = " | ".join(p for p in contact_parts if p)
    summary = tailored_resume.get("tailored_summary", tailored_resume.get("summary", ""))
    
    hydrated_bullets = tailored_resume.get("hydrated_bullets", [])
    work_histories = {w["id"]: w for w in kb.get("work_history", [])}
    
    projects = {p["id"]: p for p in kb.get("projects", [])}
    gen_proj_bullets = tailored_resume.get("generated_project_bullets", [])

    # The tech-stack line for each project (e.g. "Python · FastAPI · React 18 · Qdrant...")
    # is stored as a plain resume_bullet with parent_id == project id, NOT inside
    # generated_project_bullets. Pull those out separately so they actually render —
    # previously this line was silently missing from every generated resume.
    project_tech_lines = {}
    for b in hydrated_bullets:
        pid = b.get("parent_id")
        if pid in projects and "\u00b7" in b.get("text", ""):
            project_tech_lines[pid] = b["text"]

    skills = kb.get("skills", {})

    # me.json stores skills as a flat languages/frameworks/tools/platforms schema, but the
    # reference resume groups them into six thematic categories. Map known items into those
    # categories; anything unrecognized falls into "Additional Skills" so nothing from the KB
    # ever silently disappears just because it wasn't in this lookup table.
    _SKILL_CATEGORY_MAP = {
        "full stack": ["fastapi", "rest apis", "react 18", "javascript", "typescript",
                        "postgresql", "mongodb", "docker"],
        "ai / agentic systems": ["langgraph", "langchain", "agentic rag", "multi-agent systems",
                        "prompt engineering", "lora fine-tuning", "llamaindex", "self-rag",
                        "crag", "graphrag", "hybrid search (bm25+dense)", "crossencoder re-ranking",
                        "semantic caching", "sentence-transformers", "huggingface transformers"],
        "llms & vector dbs": ["groq", "ollama", "vllm", "openai", "aws bedrock", "medcpt",
                        "qdrant cloud", "pgvector", "pinecone"],
        "ml & evaluation": ["scikit-learn", "lightgbm", "xgboost", "tensorflow", "keras",
                        "mlflow", "ragas", "rapidfuzz", "bleu", "ndcg", "hallucination detection",
                        "drift detection", "citation accuracy", "pytest"],
        "backend & data": ["python", "duckdb", "redis", "sqlalchemy", "alembic", "pydantic",
                        "pymupdf", "docling", "pandas", "numpy"],
        "cloud & security": ["oracle oci", "aws", "gcp", "kubernetes", "hipaa", "gdpr",
                        "fda 21 cfr part 11", "aes-256-gcm", "oauth2", "jwt", "opentelemetry",
                        "prometheus", "linux", "git", "sql"],
    }

    def _grouped_skills(skills_dict):
        all_items = []
        for bucket in ("languages", "frameworks", "tools", "platforms"):
            all_items.extend(skills_dict.get(bucket, []))
        grouped = {cat: [] for cat in _SKILL_CATEGORY_MAP}
        grouped["additional skills"] = []
        seen = set()
        for item in all_items:
            if item in seen:
                continue
            seen.add(item)
            placed = False
            for cat, known in _SKILL_CATEGORY_MAP.items():
                if item.lower() in known:
                    grouped[cat].append(item)
                    placed = True
                    break
            if not placed:
                grouped["additional skills"].append(item)
        return grouped
    
    def build_flowables(font_size, bullets_to_include):
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'NameTitle',
            parent=styles['Heading1'],
            fontSize=font_size + 6,
            alignment=TA_CENTER,
            spaceAfter=4
        )
        
        contact_style = ParagraphStyle(
            'ContactInfo',
            parent=styles['Normal'],
            fontSize=font_size - 1,
            alignment=TA_CENTER,
            spaceAfter=8
        )
        
        heading_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=font_size + 2,
            spaceBefore=6,
            spaceAfter=2,
            fontName='Helvetica-Bold'
        )
        
        # We define a tab stop for right-aligned dates. Page width is 612. Margin left is 36, right is 36.
        # So width = 612 - 72 = 540.
        job_title_style = ParagraphStyle(
            'JobTitle',
            parent=styles['Heading3'],
            fontSize=font_size,
            spaceBefore=4,
            spaceAfter=2,
            fontName='Helvetica-Bold',
            tabStops=[(540, 'right')]
        )
        
        normal_style = ParagraphStyle(
            'NormalBullet',
            parent=styles['Normal'],
            fontSize=font_size,
            leading=font_size + 2,
            spaceAfter=2,
            leftIndent=15,
            firstLineIndent=-10
        )
        
        from reportlab.platypus import HRFlowable
        from reportlab.lib.colors import black
        
        def section_header(title):
            return [
                Paragraph(f"<b>{title}</b>", heading_style),
                HRFlowable(width="100%", thickness=0.5, color=black, spaceBefore=0, spaceAfter=4)
            ]
        
        elements = []
        elements.append(Paragraph(name, title_style))
        elements.append(Paragraph(contact_info, contact_style))
        
        if summary:
            elements.extend(section_header("OBJECTIVE"))
            elements.append(Paragraph(summary, ParagraphStyle('Summ', parent=styles['Normal'], fontSize=font_size, leading=font_size+2)))
        
        exp_map = {}
        # Preserve original order for display, but filter by bullets_to_include
        include_ids = {b["id"] for b in bullets_to_include if "id" in b}
        for b in hydrated_bullets:
            if "id" in b and b["id"] not in include_ids:
                continue
            pid = b.get("parent_id")
            if not pid: continue
            if pid not in exp_map:
                parent = work_histories.get(pid)
                if parent:
                    exp_map[pid] = {
                        "company": parent.get("company", ""),
                        "role": parent.get("role", ""),
                        "location": parent.get("location", ""),
                        "date": f"{parent.get('start_date','')} - {parent.get('end_date','')}",
                        "bullets": []
                    }
            if pid in exp_map:
                exp_map[pid]["bullets"].append(b["text"])
                
        if exp_map:
            elements.extend(section_header("WORK EXPERIENCE"))
            for pid, job in exp_map.items():
                # Reference format: "Role — Company, Location    Date" (role leads, not company)
                company_line = job['company']
                if job.get('location'):
                    company_line += f", {job['location']}"
                header = f"{job['role']} \u2014 {company_line}<t/><font color='gray'>{job['date']}</font>"
                elements.append(Paragraph(header, job_title_style))
                for text in job['bullets']:
                    elements.append(Paragraph(f"• {text}", normal_style))
                    
        if gen_proj_bullets:
            # Pre-filter to only valid projects
            valid_projects = []
            for pb in gen_proj_bullets:
                pid = pb.get("project_id")
                parent = projects.get(pid)
                if parent:
                    valid_projects.append((pid, parent, pb.get("bullets", [])))
            
            if valid_projects:
                elements.extend(section_header("TECHNICAL PROJECTS"))
                for pid, parent, bullets in valid_projects:
                    name_and_tagline = parent.get('name', '')
                    if parent.get('tagline'):
                        name_and_tagline += f" \u2014 {parent.get('tagline')}"
                    gh = parent.get('github', '')
                    header = f"{name_and_tagline}<t/><font color='gray'>{gh}</font>"
                    elements.append(Paragraph(header, job_title_style))
                    tech_line = project_tech_lines.get(pid)
                    if tech_line:
                        elements.append(Paragraph(tech_line, ParagraphStyle(
                            'TechLine', parent=styles['Normal'], fontSize=font_size - 1,
                            textColor='#555555', spaceAfter=2)))
                    for text in bullets:
                        elements.append(Paragraph(f"• {text}", normal_style))

        # Combined "EDUCATION & CERTIFICATIONS" section, matching the reference resume.
        education = kb.get("education", [])
        certs = kb.get("certifications", [])
        if education or certs:
            elements.extend(section_header("EDUCATION & CERTIFICATIONS"))
            for edu in education:
                header = f"{edu.get('degree', '')} \u2014 {edu.get('institution', '')}<t/><font color='gray'>{edu.get('year', '')}</font>"
                elements.append(Paragraph(header, job_title_style))
                gpa = edu.get('gpa') or edu.get('cgpa')
                if gpa:
                    elements.append(Paragraph(f"CGPA: {gpa}", normal_style))
            if certs:
                cert_lines = " \u00b7 ".join(c.get('name', '') for c in certs if c.get('name'))
                if cert_lines:
                    elements.append(Paragraph(f"<b>Certifications:</b> {cert_lines}",
                        ParagraphStyle('Certs', parent=styles['Normal'], fontSize=font_size, leading=font_size+2)))

        elements.extend(section_header("TECHNICAL SKILLS"))
        s_style = ParagraphStyle('Skills', parent=styles['Normal'], fontSize=font_size, leading=font_size+2)
        _display_names = {
            "full stack": "Full Stack",
            "ai / agentic systems": "AI / Agentic Systems",
            "llms & vector dbs": "LLMs & Vector DBs",
            "ml & evaluation": "ML & Evaluation",
            "backend & data": "Backend & Data",
            "cloud & security": "Cloud & Security",
            "additional skills": "Additional Skills",
        }
        grouped = _grouped_skills(skills)
        for key in ("full stack", "ai / agentic systems", "llms & vector dbs",
                    "ml & evaluation", "backend & data", "cloud & security", "additional skills"):
            items = grouped.get(key, [])
            if items:
                elements.append(Paragraph(f"<b>{_display_names[key]}:</b> {', '.join(items)}", s_style))

        return elements

    # We preserve the original chronological ordering provided by tailor.py.
    # When we need to truncate, we drop the oldest bullets (from the bottom).
    current_bullets = [b for b in hydrated_bullets if "id" in b]
    
    class FitDocTemplate(SimpleDocTemplate):
        def handle_pageBegin(self):
            super().handle_pageBegin()
            if self.page > 1:
                raise OverflowError("Too many pages")
                
    class StopFitting(Exception): pass
    
    best_font_size = None
    try:
        while True:
            for font_size in [11, 10.5, 10, 9.5, 9]:
                elements = build_flowables(font_size, current_bullets)
                
                # Test build in memory
                from io import BytesIO
                dummy = BytesIO()
                test_doc = FitDocTemplate(dummy, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
                
                try:
                    test_doc.build(elements)
                    # If we get here, it fit on one page!
                    best_font_size = font_size
                    print(f"[Formatter] Fit successfully at font size {font_size}")
                    raise StopFitting()
                except OverflowError:
                    print(f"[Debug] Overflowed at font size {font_size}")
                    continue
            
            # If we exhausted font sizes and still didn't fit, drop the oldest bullet (from the end of the list)
            if current_bullets:
                dropped = current_bullets.pop(-1)
                print(f"[Formatter] PDF Content too long. Truncating oldest bullet to save space: {dropped.get('id')}")
            else:
                best_font_size = 9
                raise StopFitting()
                
    except StopFitting:
        pass
        
    if best_font_size:
        try:
            final_elements = build_flowables(best_font_size, current_bullets)
            author_name = personal.get("name", "")
            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=letter,
                rightMargin=36,
                leftMargin=36,
                topMargin=36,
                bottomMargin=36,
                author=author_name,
                title=f"{author_name} Resume" if author_name else "Resume",
                creator="",
                producer=""
            )
            doc.build(final_elements)
            return True
        except Exception as e:
            print(f"[Formatter] PDF Build Error: {e}")
            return False
    return False